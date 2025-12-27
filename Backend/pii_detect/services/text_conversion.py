from docx import Document
import logging
from fastapi import HTTPException

logger = logging.getLogger("pii_detect.text_extract")

def extract_text_from_docx(file_path: str) -> str:
    try:
        doc = Document(file_path)
        text_parts = [para.text for para in doc.paragraphs if para.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            text_parts.append(para.text)

        return "\n".join(text_parts)
    except Exception as e:
        logger.exception("Text extraction from DOCX failed.")
        raise HTTPException(status_code=500, detail="Failed to extract text from .docx file.")
