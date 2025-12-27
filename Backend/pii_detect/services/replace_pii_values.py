from __future__ import annotations
from .redact_textboxes_only import redact_text_in_textboxes
from difflib import SequenceMatcher
from pathlib import Path
from typing import List, Dict, Tuple
from docx import Document


MASK_CHAR = "X"

def build_buffer(path: str | Path) -> Tuple[str, List[Tuple], Document]:
    doc = Document(path)
    buf: List[str] = []
    run_map: List[Tuple] = []

    def feed(paragraph, *, in_cell: bool):
        for run in paragraph.runs:
            pos = len(buf)
            run_map.append((run, pos, len(run.text)))
            buf.extend(run.text)
        if not in_cell:
            buf.append("\n")

    for p in doc.paragraphs:
        feed(p, in_cell=False)

    for tbl in doc.tables:
        for row in tbl.rows:
            for ci, cell in enumerate(row.cells):
                for p in cell.paragraphs:
                    feed(p, in_cell=True)
                if ci < len(row.cells) - 1:
                    buf.append("\t")
            buf.append("\n")
    for section in doc.sections:
        for hdr in [section.header, section.footer]:
            for p in hdr.paragraphs:
                feed(p, in_cell=False)


    return "".join(buf), run_map, doc


def convert_pii_in_docx_span(
        src: str | Path,
        spans: List[Dict],
        dst: str | Path,
        text_and_map: Tuple[str, List[Tuple]] = None
) -> None:
    """
    Mask substrings in a DOCX using start/end positions with fuzzy matching support.
    Handles cross-run masking, field code issues, and ensures clean overwriting.
    """
    if text_and_map is None:
        from .replace_pii_values import build_buffer
        text, run_map, doc = build_buffer(src)
    else:
        text, run_map, doc = text_and_map

    buf = list(text)
    masked = set()

    # ✅ Apply masking to the character buffer
    for s in spans:
        start = s.get("start")
        end = s.get("end")
        val = str(s.get("value", ""))

        if not val or start is None or end is None:
            continue

        found = text[start:end]
        score = SequenceMatcher(None, found, val).ratio()

        if score >= 0.9:
            print(f"✅ Fuzzy match OK: '{val}' ~ '{found}' at {start}-{end} (score={score:.2f})")
            for i in range(start, end):
                if i not in masked:
                    buf[i] = MASK_CHAR
                    masked.add(i)
        else:
            print(f"⚠️ Skipped: '{val}' did not match buffer slice '{found}' (score={score:.2f})")

    masked_text = "".join(buf)

    # ✅ Update each run with masked content
    for run, pos, ln in run_map:
        if ln == 0 or run._r.xpath(".//w:drawing"):
            continue

        original = run.text
        masked_segment = masked_text[pos: pos + ln]

        if original and original != masked_segment:
            print(f"🔧 Updating run: '{original}' → '{masked_segment}'")

            # ⚠️ Clear any child XML nodes like field codes or bookmarks
            for child in run._r.xpath("*"):
                run._r.remove(child)

            # ✅ Set the clean masked text
            run.text = masked_segment

    doc.save(dst)
    print("✅ Span-based redacted file saved to:", dst)
    pii_values = [s["value"] for s in spans if isinstance(s.get("value"), str)]
    if pii_values:
        redact_text_in_textboxes(dst, pii_values)